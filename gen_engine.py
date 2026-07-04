#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""单词练习生成引擎
读取词表，按题型规则生成题目与答案 Word 文档，输出到 成品/ 下。
"""
import os, re, random, json, time, urllib.request
from concurrent.futures import ThreadPoolExecutor, as_completed
import docx
from docx.shared import Pt, RGBColor, Mm
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from nltk.corpus import wordnet as wn
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas

random.seed(20220630)  # 固定随机种子保证可复现

# ============ 全局字体参数 ============
STANDARD_FONT = "Times New Roman"
CN_FONT = "宋体"
TITLE_PT = 16        # 题型总标题
QHEAD_PT = 14        # 每组题号标题 (第X组)
GROUP_SUB_PT = 12    # 题型说明小标题
BODY_PT = 12         # 正文 12pt

ANSWER_TITLE_PT = 18  # 答案文档大标题
ANSWER_GROUP_PT = 16  # 答案每组标题
TRANSLATION_NORMAL_PT = 10.5
TRANSLATION_ANSWER_TITLE_PT = 18
TRANSLATION_ANSWER_GROUP_PT = 16
TRANSLATION_ANSWER_LINE_PT = 16
TRANSLATION_HEADER_PT = 18

TRANSLATION_PDF_FONT = "STSong-Light"
TRANSLATION_PDF_HEADER_FONT = "Times-Bold"
TRANSLATION_PDF_PAGE_SIZE = A4
TRANSLATION_PDF_LEFT_PT = 90
TRANSLATION_PDF_RIGHT_PT = 90
TRANSLATION_PDF_TOP_PT = 72
TRANSLATION_PDF_BOTTOM_PT = 72
TRANSLATION_LAYOUT_CANDIDATES = [
    {"cols": 4, "font_size": 10.5, "spacing_pt": 10},
    {"cols": 3, "font_size": 10.5, "spacing_pt": 12},
    {"cols": 4, "font_size": 10.0, "spacing_pt": 10},
    {"cols": 3, "font_size": 10.0, "spacing_pt": 12},
    {"cols": 4, "font_size": 9.5, "spacing_pt": 10},
    {"cols": 3, "font_size": 9.5, "spacing_pt": 12},
    {"cols": 3, "font_size": 9.0, "spacing_pt": 12},
    {"cols": 3, "font_size": 8.5, "spacing_pt": 12},
]

PDF_FONT_READY = False

DISPLAY_WORD_OVERRIDES = {
    "analyse/ze": "analyze",
    "apologize/se": "apologize",
    "grey/ay": "gray",
    "organise/ze": "organize",
    "realise/ze": "realize",
    "recognise/ze": "recognize",
}

def _set_run_font(run, name=STANDARD_FONT, size=BODY_PT, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    # 中文字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)

def add_para(doc, text="", size=BODY_PT, bold=False, align=None, font=STANDARD_FONT,
             space_before=None, space_after=None, line=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    if space_before is not None: pf.space_before = Pt(space_before)
    if space_after is not None: pf.space_after = Pt(space_after)
    if line is not None: pf.line_spacing = line
    if text:
        r = p.add_run(text)
        _set_run_font(r, name=font, size=size, bold=bold)
    return p

def _set_doc_normal_style(doc, font_name=STANDARD_FONT, size_pt=BODY_PT):
    st = doc.styles['Normal']
    st.font.name = font_name
    st.font.size = Pt(size_pt)
    rPr = st.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

def new_doc(page_size="letter"):
    doc = docx.Document()
    # Normal 样式设默认字体
    _set_doc_normal_style(doc, STANDARD_FONT, BODY_PT)
    # 页边距与参考文档一致 (1英寸左右)
    for sec in doc.sections:
        if page_size == "a4":
            sec.page_width = Mm(210)
            sec.page_height = Mm(297)
        sec.top_margin = Pt(72); sec.bottom_margin = Pt(72)
        sec.left_margin = Pt(90); sec.right_margin = Pt(90)
    return doc

def set_section_columns(section, num_cols, spacing_pt=12):
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), str(num_cols))
    cols.set(qn('w:space'), str(int(spacing_pt * 20)))
    cols.set(qn('w:equalWidth'), '1')

def set_paragraph_keep(paragraph, keep_next=False, keep_lines=True):
    pPr = paragraph._element.get_or_add_pPr()
    for tag, enabled in (('w:keepNext', keep_next), ('w:keepLines', keep_lines)):
        node = pPr.find(qn(tag))
        if enabled:
            if node is None:
                node = OxmlElement(tag)
                pPr.append(node)
        elif node is not None:
            pPr.remove(node)

def set_section_header(section, text, font=STANDARD_FONT, size=TRANSLATION_HEADER_PT):
    header = section.header
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(str(text))
    _set_run_font(run, name=font, size=size, bold=True)

# ============ 词表与 WordNet 数据层 ============
def display_word(w):
    """用于题面显示的英文词条：去掉括号及其中说明，修掉明显空白问题。"""
    s = str(w).replace('\xa0', ' ').strip()
    s = s.replace("’", "'").replace("‘", "'").replace("…", "...")
    s = re.sub(r'\s+', ' ', s)
    s = re.sub(r'^[*•·]+', '', s)
    s = re.sub(r'\s*=\s*.*$', '', s)
    s = DISPLAY_WORD_OVERRIDES.get(s, s)
    s = re.sub(r'\([^)]*\)', '', s)
    s = re.sub(r'\s*\(.*$', '', s)
    s = re.sub(r"\s*'\s*", "'", s)
    s = re.sub(r'\.{3,}', ' ', s)
    s = re.sub(r'[?!]+$', '', s)
    s = re.sub(r'\s+', ' ', s).strip()
    s = s.strip(" .")
    s = re.sub(r'\s+\b(?:v|n|adj|adv|vt|vi)\.$', '', s, flags=re.IGNORECASE)
    if s.lower() == "ai":
        return "AI"
    return s

def clean_word(w):
    """取词形主干：去括号注释、小写；保留 pay phone / space station 等短语空格。"""
    s = display_word(w).lower()
    s = re.sub(r'\([^)]*\)', '', s)
    s = s.replace("’", "'").replace("…", " ")
    s = re.sub(r'[^a-z0-9\' -]+', ' ', s)
    s = re.sub(r'\s+', ' ', s).strip()
    return s

# 仅用于判断同义/反义词是否"常见可用"：纯小写字母、无连字符/下划线、长度3~12
COMMON_RE = re.compile(r'^[a-z]{3,12}$')
def is_common(w):
    return bool(COMMON_RE.match(str(w)))

def commonness(lemma_name):
    """用'该词自身的 synset 数量'近似词频/常见度"""
    try:
        return len(wn.synsets(lemma_name))
    except Exception:
        return 0

# 同义词右项至少要达到的常见度阈值(自身synset数), 用于过滤生僻 lemma
# 同时用一个小型"中小学常见词"白名单做最终把关
MIN_COMMON_SYN = 4
MIN_COMMON_ANT = 3

def _load_common_set():
    """构建'常见英文词'集合: 收集所有 synset 数 >= MIN_COMMON_SYN 的 lemma 名(纯小写字母)。
    这些词天然是英语中多义、常用的词。一次性构建约需 5~10 秒。"""
    s = set()
    for syn in wn.all_synsets():
        for l in syn.lemmas():
            nm = l.name()
            if is_common(nm) and commonness(nm) >= MIN_COMMON_SYN:
                s.add(nm)
    return s

COMMON_SET = None  # 惰性初始化
def is_common_word(w):
    global COMMON_SET
    if COMMON_SET is None:
        print("  [初始化常见词集合...]", flush=True)
        COMMON_SET = _load_common_set()
        print(f"  [常见词集合大小: {len(COMMON_SET)}]", flush=True)
    return w in COMMON_SET

def clean_cn(cn):
    """把中文释义规整成可显示文本。"""
    cn = str(cn).replace('\xa0', ' ').strip()
    cn = re.sub(r'\s+', ' ', cn)
    return cn

POS_LABEL_RE = re.compile(
    r'(?<![A-Za-z])(?:vt|vi|v|n|adj|adv|a|ad|pron|prep|conj|num|int|interj|det|aux|pl)\.\s*',
    re.IGNORECASE,
)

def plain_cn(cn):
    """汉译英/英译汉题面使用的中文释义：去掉词性标记，贴近参考模板。"""
    s = clean_cn(cn)
    s = s.replace('＆', '&')
    s = POS_LABEL_RE.sub('', s)
    s = s.replace(';', '；').replace(',', '，')
    s = re.sub(r'([\u4e00-\u9fff]{2,6})\1', r'\1', s)
    s = re.sub(r'\s+', ' ', s).strip(' &;；,，')
    return s or clean_cn(cn)

def pos_priority(cn):
    """根据中文释义中的词性标记优先选择 WordNet 词性，减少 definition 错义。"""
    text = clean_cn(cn).lower()
    order = []
    mapping = {
        "n": wn.NOUN,
        "v": wn.VERB,
        "vt": wn.VERB,
        "vi": wn.VERB,
        "adj": wn.ADJ,
        "a": wn.ADJ,
        "adv": wn.ADV,
        "ad": wn.ADV,
    }
    for m in re.finditer(r'\b(vt|vi|adj|adv|ad|v|n|a)\.', text):
        pos = mapping.get(m.group(1))
        if pos and pos not in order:
            order.append(pos)
    for pos in (wn.NOUN, wn.VERB, wn.ADJ, wn.ADV):
        if pos not in order:
            order.append(pos)
    return order

def wordnet_keys(en):
    base = clean_word(en)
    if not base:
        return []
    return [base.replace(' ', '_')]

def spelling_core(en):
    """乱序/缺字母只使用单个纯字母词，短语不再被粘成假单词。"""
    base = clean_word(en)
    return base if base and ' ' not in base and base.isalpha() else ""

def _letters_only(s):
    return re.sub(r'[^a-z]', '', str(s).lower())

def _simple_stem(s):
    s = _letters_only(s)
    for suf in (
        "ingly", "edly", "ation", "itions", "ness", "ment", "tion", "sion",
        "ity", "ist", "ism", "ive", "ous", "able", "ible", "ful", "less",
        "ing", "ed", "ly", "er", "est", "al", "s",
    ):
        if len(s) > len(suf) + 3 and s.endswith(suf):
            return s[:-len(suf)]
    return s

def is_morph_variant(a, b):
    aa, bb = _letters_only(a), _letters_only(b)
    if not aa or not bb or aa == bb:
        return True
    if _simple_stem(aa) == _simple_stem(bb):
        return True
    if min(len(aa), len(bb)) >= 4 and abs(len(aa) - len(bb)) <= 5:
        return aa.startswith(bb) or bb.startswith(aa)
    return False

def normalize_relation_term(term):
    term = str(term).replace('_', ' ').replace('-', ' ').lower()
    term = re.sub(r'[^a-z ]+', ' ', term)
    term = re.sub(r'\s+', ' ', term).strip()
    return term

def is_relation_term(term, head):
    term = normalize_relation_term(term)
    head = normalize_relation_term(head)
    if not term or not re.match(r'^[a-z]+(?: [a-z]+)?$', term):
        return False
    compact = term.replace(' ', '')
    if len(compact) < 3 or len(compact) > 18:
        return False
    if is_morph_variant(term, head):
        return False
    return True

BAD_RELATION_TERMS = {
    "attic", "bean counter", "boodle", "bonce", "cabbage", "can", "chisel",
    "chouse", "clams", "dinero", "dome", "encephalon", "forrader", "gelt",
    "humourous", "jakes", "john", "lav", "lavatory", "loot", "maestro",
    "method acting", "monish", "noggin", "noodle", "pot", "potty", "privy",
    "scratch", "staff of life", "stool", "sugar", "tierce", "twelvemonth",
    "wampum", "woman chaser",
}

BAD_RELATION_HEADS = {
    "bean", "bread", "season", "wolf", "year",
}

BAD_ANTONYM_HEADS = {
    "peel",
}

ALLOWED_SIMILAR_SYNONYMS = {
    "dormitory": {"dorm"},
    "electric": {"electrical"},
    "rainfall": {"rain"},
    "percentage": {"percent"},
    "spring": {"springtime"},
}

CURATED_SYNONYMS = {
    "ahead": ["forward"],
    "always": ["forever"],
    "brain": ["mind"],
    "boy": ["lad"],
    "call back": ["recall"],
    "cheat": ["deceive"],
    "class": ["grade"],
    "community": ["society"],
    "countryside": ["village"],
    "dead": ["lifeless"],
    "eat out": ["dine out"],
    "educate": ["teach"],
    "energy": ["power"],
    "go along": ["proceed"],
    "handbag": ["purse"],
    "home": ["house"],
    "humorous": ["funny"],
    "keyboard": ["keypad"],
    "kill": ["murder"],
    "less than": ["under"],
    "master": ["expert"],
    "method": ["way"],
    "offer": ["provide"],
    "pay phone": ["public phone"],
    "peel": ["pare"],
    "quick": ["fast"],
    "receive": ["get"],
    "relax": ["rest"],
    "return": ["come back"],
    "share": ["divide"],
    "space station": ["space laboratory"],
    "teacher": ["instructor"],
    "toilet": ["restroom"],
    "touch": ["feel"],
    "warn": ["alert"],
    "wide": ["broad"],
    "win": ["succeed"],
}

CURATED_ANTONYMS = {
    "ahead": ["back"],
    "always": ["never"],
    "boy": ["girl"],
    "dead": ["alive"],
    "eat out": ["eat in"],
    "go along": ["stop"],
    "lose": ["win"],
    "quick": ["slow"],
    "relax": ["tense"],
    "uncle": ["aunt"],
    "wide": ["narrow"],
    "win": ["lose"],
}

BAD_REQUIRED_SYNONYMS = {
    "bang", "beam", "canvass", "canvas", "charge", "company", "conk", "fixture",
    "fixing", "flush", "invidia", "intumesce", "mending", "sanction", "tumefy",
    "tumesce",
}

GENERATION_SYNONYM_OVERRIDES = {
    "abandon": "leave",
    "abuse": "harm",
    "accelerate": "speed",
    "acknowledge": "admit",
    "accompany": "follow",
    "accuse": "blame",
    "achieve": "reach",
    "appreciate": "value",
    "approve": "accept",
    "assemble": "gather",
    "assist": "help",
    "astonish": "surprise",
    "award": "prize",
    "beg": "ask",
    "broadcast": "send",
    "capture": "catch",
    "cherish": "value",
    "click": "press",
    "conduct": "carry out",
    "confuse": "puzzle",
    "contain": "hold",
    "continue": "go on",
    "criticise": "blame",
    "declare": "say",
    "decrease": "reduce",
    "delay": "postpone",
    "depend": "rely",
    "develop": "grow",
    "devote": "give",
    "dive": "jump",
    "drown": "sink",
    "enhance": "improve",
    "engineer": "design",
    "evolve": "develop",
    "exact": "precise",
    "examine": "study",
    "exceed": "pass",
    "exclaim": "cry",
    "expand": "grow",
    "expose": "reveal",
    "factor": "cause",
    "fasten": "tie",
    "float": "drift",
    "govern": "rule",
    "hesitate": "pause",
    "import": "bring in",
    "include": "contain",
    "indicate": "show",
    "infer": "guess",
    "injure": "hurt",
    "inspire": "encourage",
    "insert": "put in",
    "lean": "tilt",
    "measure": "size",
    "mention": "say",
    "notice": "see",
    "observe": "notice",
    "occur": "happen",
    "overlook": "miss",
    "persist": "continue",
    "polish": "shine",
    "position": "place",
    "pour": "flow",
    "postpone": "delay",
    "pressure": "stress",
    "produce": "make",
    "react": "respond",
    "realize": "understand",
    "recognize": "spot",
    "recommend": "suggest",
    "record": "write down",
    "reflect": "show",
    "regulation": "rule",
    "repair": "fix",
    "resemble": "look like",
    "resist": "fight",
    "respond": "answer",
    "roll": "turn",
    "result": "outcome",
    "reveal": "show",
    "root": "base",
    "scan": "examine",
    "search": "look for",
    "scold": "criticize",
    "sculpture": "statue",
    "shave": "trim",
    "simulate": "copy",
    "speed": "hurry",
    "stare": "gaze",
    "suspect": "think",
    "swell": "grow",
    "sweep": "brush",
    "tap": "touch",
    "tool": "instrument",
    "thrill": "excite",
    "translate": "interpret",
    "yield": "produce",
}

DISALLOWED_TYPE3_HEADS = {
    "arrest",
    "board",
    "envy",
    "faint",
    "heat",
}

def curated_terms(mapping, head):
    return [normalize_relation_term(x) for x in mapping.get(normalize_relation_term(head), [])]

def allowed_similar_synonym(word, syn):
    return normalize_relation_term(syn) in ALLOWED_SIMILAR_SYNONYMS.get(normalize_relation_term(word), set())

def unique_distractors(pool, correct, n, key=lambda x: x):
    """取不与正确项重复、彼此也不重复的干扰项。"""
    correct_key = key(correct)
    seen = {correct_key}
    out = []
    cand = list(pool)
    random.shuffle(cand)
    for item in cand:
        k = key(item)
        if k in seen:
            continue
        seen.add(k)
        out.append(item)
        if len(out) >= n:
            break
    return out

def preferred_synonym_for_word(wd_item):
    if not wd_item:
        return ""
    for syn in list(getattr(wd_item, "syns_common", [])) + list(getattr(wd_item, "syns", [])):
        syn = normalize_relation_term(syn)
        if syn:
            return syn
    return ""

# ============ LLM 题面材料缓存 ============
LLM_CACHE_FILE = os.path.join(os.path.dirname(__file__), "llm_sentence_cache.json")
LLM_CACHE = None
BASIC_MATERIAL_FIELDS = ("cloze_sentence", "tf_true", "tf_false")
SYNONYM_MATERIAL_FIELDS = ("synonym", "synonym_original", "synonym_rewrite_blank")

def material_key(en, cn):
    return f"{clean_word(en)}||{plain_cn(cn)}"

def load_llm_cache():
    global LLM_CACHE
    if LLM_CACHE is None:
        if os.path.exists(LLM_CACHE_FILE):
            with open(LLM_CACHE_FILE, "r", encoding="utf-8") as f:
                LLM_CACHE = json.load(f)
        else:
            LLM_CACHE = {}
    return LLM_CACHE

def save_llm_cache():
    if LLM_CACHE is not None:
        with open(LLM_CACHE_FILE, "w", encoding="utf-8") as f:
            json.dump(LLM_CACHE, f, ensure_ascii=False, indent=2)

def has_material_fields(entry, require_synonym=False, required_synonym=""):
    if not isinstance(entry, dict) or entry.get("_source") != "llm":
        return False
    needed = SYNONYM_MATERIAL_FIELDS if require_synonym else BASIC_MATERIAL_FIELDS
    if not all(str(entry.get(field, "")).strip() for field in needed):
        return False
    if require_synonym and required_synonym:
        cached_syn = normalize_relation_term(entry.get("_required_synonym", "") or entry.get("synonym", ""))
        if cached_syn != normalize_relation_term(required_synonym):
            return False
    return True

def strip_code_fence(text):
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r'^```(?:json)?\s*', '', text)
        text = re.sub(r'\s*```$', '', text)
    return text.strip()

def replace_answer_with_blank(sentence, answer):
    sentence = str(sentence).strip()
    answer = str(answer).strip()
    if not sentence or not answer:
        return ""
    pattern = re.escape(answer)
    pattern = pattern.replace(r'\ ', r'\s+')
    replaced, n = re.subn(
        rf'(?<![A-Za-z]){pattern}(?![A-Za-z])',
        '______',
        sentence,
        count=1,
        flags=re.IGNORECASE,
    )
    return replaced if n == 1 else ""

def sanitize_material(raw, en, cn, wd_item=None, require_synonym=False):
    if not isinstance(raw, dict):
        raise ValueError("material is not an object")
    word = display_word(en)
    required_synonym = preferred_synonym_for_word(wd_item)
    out = {}
    if require_synonym:
        rewrite_full = str(raw.get("synonym_rewrite_full") or raw.get("synonym_rewrite_blank") or "").strip()
        if not rewrite_full:
            raise ValueError("missing synonym_rewrite_full")
        for field in ("synonym", "synonym_original"):
            val = str(raw.get(field, "")).strip()
            if not val:
                raise ValueError(f"missing {field}")
            val = re.sub(r'_{3,}', '______', val)
            out[field] = val
        syn = normalize_relation_term(required_synonym or out["synonym"])
        if not syn or len(syn) > 32:
            raise ValueError("bad synonym")
        out["synonym"] = syn
        out["synonym_rewrite_blank"] = replace_answer_with_blank(rewrite_full, syn)
        if out["synonym_rewrite_blank"].count("______") != 1:
            raise ValueError("synonym_rewrite_full must contain the synonym exactly once")
        out["_required_synonym"] = syn
    else:
        for field in ("tf_true", "tf_false"):
            val = str(raw.get(field, "")).strip()
            if not val:
                raise ValueError(f"missing {field}")
            val = re.sub(r'_{3,}', '______', val)
            out[field] = val
        cloze_full = str(raw.get("cloze_full_sentence") or raw.get("cloze_sentence") or "").strip()
        if not cloze_full:
            raise ValueError("missing cloze_full_sentence")
        out["cloze_sentence"] = replace_answer_with_blank(cloze_full, word)
        if out["cloze_sentence"].count("______") != 1:
            raise ValueError("cloze_full_sentence must contain the target word exactly once")
    out["_source"] = "llm"
    return out

def call_llm_for_materials(items, require_synonym=False):
    api_key = os.getenv("VIVI_LLM_API_KEY", "")
    if not api_key:
        raise RuntimeError("缺少 VIVI_LLM_API_KEY，不能生成 LLM 题面")
    base_url = os.getenv("VIVI_LLM_BASE_URL", "https://api.070320.xyz").rstrip("/")
    if not base_url.startswith("http"):
        base_url = "https://" + base_url
    model = os.getenv("VIVI_LLM_MODEL", "deepseek-v4-flash")
    url = base_url + "/v1/chat/completions"
    single_item_strict = ""
    temperature = 0.2
    if len(items) == 1:
        if require_synonym:
            single_item_strict = (
                " Single-item strict mode: choose the final exact text for synonym first. "
                "If grammar needs an inflected form, put that final inflected form in synonym. "
                "Then copy that exact synonym text unchanged into synonym_rewrite_full once and only once. "
                "Also make sure cloze_full_sentence contains the exact word field once and only once."
            )
        else:
            single_item_strict = (
                " Single-item strict mode: make sure cloze_full_sentence contains the exact word field once and only once."
            )
        temperature = 0.0
    requested_fields = "cloze_full_sentence, tf_true, tf_false"
    synonym_rules = ""
    cloze_rules = (
        "Rules for cloze_full_sentence: it is for a multiple-choice item whose option is the vocabulary headword, "
        "so it must contain the exact input word string once, with no inflection, plural, tense change, or added suffix. "
        "Write the sentence so that exact form is grammatical; for verbs, use patterns like can/will/to + word when needed. "
        "Do not use blanks in cloze_full_sentence. It must give enough context to choose that exact word after the word is removed. "
    )
    tf_rules = (
        "tf_true and tf_false must be complete standalone English sentences using the input word; "
        "tf_true must be clearly true, tf_false clearly false. Avoid slang, rare senses, and definitions as wording. "
    )
    example_input = "[{\"id\":\"invent||发明\",\"word\":\"invent\",\"meaning\":\"发明\"}]"
    example_output = (
        "[{\"id\":\"invent||发明\","
        "\"cloze_full_sentence\":\"Young people can invent helpful tools for daily life.\","
        "\"tf_true\":\"People invent new things to solve problems.\","
        "\"tf_false\":\"People invent the sun every morning.\"}]"
    )
    payload_items = []
    if require_synonym:
        requested_fields = "synonym, synonym_original, synonym_rewrite_full"
        synonym_rules = (
            " Rules for synonym: it is the exact answer shown in the options for synonym replacement. "
            "It may be inflected if grammar requires it, but synonym_rewrite_full must contain that exact synonym string once. "
            "If an input item provides required_synonym, you must use that exact text as synonym and also place that exact text in synonym_rewrite_full once. "
            "synonym_original must be a complete sentence using the input word naturally. "
            "synonym_rewrite_full must be a very similar complete sentence containing synonym exactly once; do not use blanks in it. "
        )
        example_input = (
            "[{\"id\":\"invent||发明\",\"word\":\"invent\",\"meaning\":\"发明\",\"required_synonym\":\"create\"}]"
        )
        example_output = (
            "[{\"id\":\"invent||发明\","
            "\"synonym\":\"create\","
            "\"synonym_original\":\"Young people often invent helpful tools for daily life.\","
            "\"synonym_rewrite_full\":\"Young people often create helpful tools for daily life.\"}]"
        )
    for it in items:
        row = {"id": it["id"], "word": it["word"], "meaning": it["meaning"]}
        if require_synonym and it.get("required_synonym"):
            row["required_synonym"] = it["required_synonym"]
        payload_items.append(row)
    system = (
        "You create English vocabulary worksheet items for Chinese students. "
        "Return valid JSON only, as an array. For each input item, keep the same id and produce: "
        f"{requested_fields}. "
        + ("" if require_synonym else cloze_rules)
        + synonym_rules +
        ("" if require_synonym else tf_rules) +
        "Important: never write a clue for a different answer. For example, if the word is 'invent', write a sentence like "
        "'Scientists invent useful machines', not 'Alexander Graham Bell invented the telephone'. "
        f"Example input: {example_input} "
        f"Example output: {example_output}"
        + single_item_strict
    )
    user = json.dumps(payload_items, ensure_ascii=False)
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "temperature": temperature,
        "max_tokens": max(2048, min(32000, 260 * len(items))),
    }
    req = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Authorization": "Bearer " + api_key, "Content-Type": "application/json"},
    )
    for attempt in range(6):
        try:
            with urllib.request.urlopen(req, timeout=600) as resp:
                data = json.loads(resp.read())
            content = data["choices"][0]["message"]["content"]
            parsed = json.loads(strip_code_fence(content))
            if isinstance(parsed, dict) and "items" in parsed:
                parsed = parsed["items"]
            if not isinstance(parsed, list):
                raise ValueError("LLM response is not a list")
            return {str(obj.get("id")): obj for obj in parsed if isinstance(obj, dict)}
        except Exception as ex:
            print(f"  [LLM题面] 重试 {attempt+1}: {str(ex)[:120]}", flush=True)
            time.sleep(2 * (attempt + 1))
    raise RuntimeError("LLM 题面生成多次失败")

def call_llm_for_synonym_repair(word, meaning, synonym, original_sentence):
    api_key = os.getenv("VIVI_LLM_API_KEY", "")
    if not api_key:
        raise RuntimeError("缺少 VIVI_LLM_API_KEY，不能修复 LLM 同义替换")
    base_url = os.getenv("VIVI_LLM_BASE_URL", "https://api.070320.xyz").rstrip("/")
    if not base_url.startswith("http"):
        base_url = "https://" + base_url
    model = os.getenv("VIVI_LLM_MODEL", "deepseek-v4-flash")
    url = base_url + "/v1/chat/completions"
    system = (
        "You repair one English synonym replacement item. "
        "Return valid JSON only as one object with synonym_original and synonym_rewrite_full. "
        "Keep the two sentences very similar in meaning and structure. "
        "synonym_rewrite_full must contain the exact given synonym string once and only once. "
        "Do not use blanks."
    )
    user = json.dumps({
        "word": word,
        "meaning": meaning,
        "synonym": synonym,
        "synonym_original": original_sentence,
    }, ensure_ascii=False)
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "temperature": 0.0,
        "max_tokens": 512,
    }
    req = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Authorization": "Bearer " + api_key, "Content-Type": "application/json"},
    )
    for attempt in range(4):
        try:
            with urllib.request.urlopen(req, timeout=300) as resp:
                data = json.loads(resp.read())
            content = data["choices"][0]["message"]["content"]
            parsed = json.loads(strip_code_fence(content))
            if not isinstance(parsed, dict):
                raise ValueError("repair response is not an object")
            return parsed
        except Exception as ex:
            print(f"  [LLM题面修复] 重试 {attempt+1}: {str(ex)[:120]}", flush=True)
            time.sleep(2 * (attempt + 1))
    raise RuntimeError("LLM 同义替换修复多次失败")

def resolve_material_batch(batch_items, wd=None, batch_label="", require_synonym=False):
    """拉取并校验一批 LLM 题面；顽固失败的批次自动拆成单词级重试。"""
    pending = list(batch_items)
    resolved = {}
    tries = 0
    prefix = f"  [LLM题面][{batch_label}] " if batch_label else "  [LLM题面] "
    while pending:
        tries += 1
        print(f"{prefix}生成 {len(pending)} 条，尝试{tries}", flush=True)
        got = call_llm_for_materials(pending, require_synonym=require_synonym)
        next_pending = []
        for item in pending:
            key = item["id"]
            wd_item = wd.get((item["en"], item["cn"])) if wd else None
            raw_item = got.get(key)
            try:
                resolved[key] = sanitize_material(
                    raw_item, item["en"], item["cn"], wd_item, require_synonym=require_synonym
                )
            except Exception as ex:
                if require_synonym and len(pending) == 1 and tries >= 3:
                    expected_syn = preferred_synonym_for_word(wd_item)
                    original_sentence = str((raw_item or {}).get("synonym_original", "")).strip()
                    if expected_syn and original_sentence:
                        try:
                            repaired = call_llm_for_synonym_repair(
                                display_word(item["en"]),
                                plain_cn(item["cn"]),
                                expected_syn,
                                original_sentence,
                            )
                            merged = dict(raw_item or {})
                            merged["synonym"] = expected_syn
                            merged["synonym_original"] = str(repaired.get("synonym_original") or original_sentence).strip()
                            merged["synonym_rewrite_full"] = str(repaired.get("synonym_rewrite_full") or "").strip()
                            resolved[key] = sanitize_material(
                                merged, item["en"], item["cn"], wd_item, require_synonym=True
                            )
                            continue
                        except Exception as repair_ex:
                            print(f"{prefix}修复失败 {display_word(item['en'])}: {str(repair_ex)[:80]}", flush=True)
                print(f"{prefix}校验失败 {display_word(item['en'])}: {str(ex)[:80]}", flush=True)
                next_pending.append(item)
        if not next_pending:
            break
        if tries >= 4 and len(next_pending) > 1:
            split_total = len(next_pending)
            for idx, item in enumerate(next_pending, 1):
                single_label = f"{batch_label}|{idx}/{split_total}:{display_word(item['en'])}" if batch_label else display_word(item["en"])
                resolved.update(resolve_material_batch([item], wd, single_label, require_synonym=require_synonym))
            return resolved
        if tries >= 8:
            raise RuntimeError(f"LLM 题面校验失败次数过多: {batch_label or len(batch_items)}")
        pending = next_pending
    return resolved

def ensure_llm_materials(entries, wd=None, require_synonym=False):
    """确保指定词条的题面材料存在；entries 为 [(en, cn), ...]。"""
    cache = load_llm_cache()
    missing = []
    seen = set()
    for en, cn in entries:
        key = material_key(en, cn)
        if key in seen:
            continue
        seen.add(key)
        wd_item = wd.get((en, cn)) if wd else None
        required_synonym = preferred_synonym_for_word(wd_item) if require_synonym else ""
        if not has_material_fields(
            cache.get(key), require_synonym=require_synonym, required_synonym=required_synonym
        ):
            missing.append({
                "id": key,
                "word": display_word(en),
                "meaning": plain_cn(cn),
                "required_synonym": required_synonym,
                "en": en,
                "cn": cn,
            })
    batch_size = int(os.getenv("VIVI_LLM_BATCH_SIZE", "60"))
    concurrency = max(1, int(os.getenv("VIVI_LLM_CONCURRENCY", "5")))
    batches = [missing[i:i + batch_size] for i in range(0, len(missing), batch_size)]
    total_batches = len(batches)
    if total_batches == 0:
        return cache
    workers = min(concurrency, total_batches)
    print(f"  [LLM题面] 缺失 {len(missing)} 条，batch={batch_size}，并发={workers}", flush=True)
    if workers == 1:
        for idx, batch in enumerate(batches, 1):
            resolved = resolve_material_batch(batch, wd, f"{idx}/{total_batches}", require_synonym=require_synonym)
            for key, value in resolved.items():
                cache[key] = {**cache.get(key, {}), **value}
            save_llm_cache()
        return cache
    with ThreadPoolExecutor(max_workers=workers) as executor:
        futures = {
            executor.submit(resolve_material_batch, batch, wd, f"{idx}/{total_batches}", require_synonym): idx
            for idx, batch in enumerate(batches, 1)
        }
        errors = []
        for future in as_completed(futures):
            try:
                resolved = future.result()
            except Exception as ex:
                errors.append(str(ex))
                continue
            for key, value in resolved.items():
                cache[key] = {**cache.get(key, {}), **value}
            save_llm_cache()
        if errors:
            raise RuntimeError(errors[0])
    return cache

def get_material(en, cn, wd=None, require_synonym=False):
    cache = ensure_llm_materials([(en, cn)], wd, require_synonym=require_synonym)
    key = material_key(en, cn)
    if key not in cache:
        raise RuntimeError(f"缺少 LLM 题面: {display_word(en)}")
    return cache[key]

def has_lexical_fields(entry):
    return isinstance(entry, dict) and entry.get("_lexical_source") == "llm" and str(entry.get("definition_en", "")).strip()

def sanitize_lexical(raw, en, cn):
    if not isinstance(raw, dict):
        raise ValueError("lexical item is not an object")
    word = display_word(en)
    allow_repeat_in_definition = word.lower() in {
        "to", "for", "in", "on", "at", "by", "as", "if", "or", "of", "up", "out",
    }
    definition = re.sub(r'\s+', ' ', str(raw.get("definition_en", "")).strip())
    if not definition:
        raise ValueError("missing definition_en")
    if (not allow_repeat_in_definition) and re.search(
        rf'(?<![A-Za-z]){re.escape(word)}(?![A-Za-z])', definition, re.IGNORECASE
    ):
        raise ValueError("definition_en should not repeat the target word")
    syn = normalize_relation_term(raw.get("synonym", ""))
    ant = normalize_relation_term(raw.get("antonym", ""))
    if syn and is_morph_variant(syn, word) and not allowed_similar_synonym(word, syn):
        syn = ""
    if ant and is_morph_variant(ant, word):
        ant = ""
    return {
        "definition_en": definition,
        "synonym": syn,
        "antonym": ant,
        "_lexical_source": "llm",
    }

def call_llm_for_lexical(items):
    api_key = os.getenv("VIVI_LLM_API_KEY", "")
    if not api_key:
        raise RuntimeError("缺少 VIVI_LLM_API_KEY，不能生成 LLM 词汇关系")
    base_url = os.getenv("VIVI_LLM_BASE_URL", "https://api.070320.xyz").rstrip("/")
    if not base_url.startswith("http"):
        base_url = "https://" + base_url
    model = os.getenv("VIVI_LLM_MODEL", "deepseek-v4-flash")
    url = base_url + "/v1/chat/completions"
    system = (
        "You create vocabulary metadata for Chinese middle-school students. "
        "Return valid JSON only, as an array. For each input item, keep the same id and produce: "
        "definition_en, synonym, antonym. "
        "Rules for definition_en: short, clear English explanation, do not repeat the target word, no brackets. "
        "Rules for synonym and antonym: use very common classroom-friendly English words or short phrases, base form only, "
        "same part of speech when possible. If there is no safe, common choice, return an empty string. "
        "Avoid rare, slang, archaic, or highly technical words. "
        "Example input: [{\"id\":\"invent||发明\",\"word\":\"invent\",\"meaning\":\"发明\"}] "
        "Example output: [{\"id\":\"invent||发明\",\"definition_en\":\"to create something new\",\"synonym\":\"create\",\"antonym\":\"copy\"}]"
    )
    user = json.dumps([
        {"id": material_key(en, cn), "word": display_word(en), "meaning": plain_cn(cn)}
        for en, cn in items
    ], ensure_ascii=False)
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "temperature": 0.2,
        "max_tokens": max(2048, min(32000, 180 * len(items))),
    }
    req = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Authorization": "Bearer " + api_key, "Content-Type": "application/json"},
    )
    for attempt in range(6):
        try:
            with urllib.request.urlopen(req, timeout=600) as resp:
                data = json.loads(resp.read())
            content = data["choices"][0]["message"]["content"]
            parsed = json.loads(strip_code_fence(content))
            if isinstance(parsed, dict) and "items" in parsed:
                parsed = parsed["items"]
            if not isinstance(parsed, list):
                raise ValueError("LLM lexical response is not a list")
            return {str(obj.get("id")): obj for obj in parsed if isinstance(obj, dict)}
        except Exception as ex:
            print(f"  [LLM词汇] 重试 {attempt+1}: {str(ex)[:120]}", flush=True)
            time.sleep(2 * (attempt + 1))
    raise RuntimeError("LLM 词汇关系生成多次失败")

def resolve_lexical_batch(batch_items, batch_label=""):
    pending = list(batch_items)
    resolved = {}
    tries = 0
    prefix = f"  [LLM词汇][{batch_label}] " if batch_label else "  [LLM词汇] "
    while pending:
        tries += 1
        print(f"{prefix}生成 {len(pending)} 条，尝试{tries}", flush=True)
        got = call_llm_for_lexical([(item["en"], item["cn"]) for item in pending])
        next_pending = []
        for item in pending:
            key = item["id"]
            try:
                resolved[key] = sanitize_lexical(got.get(key), item["en"], item["cn"])
            except Exception as ex:
                print(f"{prefix}校验失败 {display_word(item['en'])}: {str(ex)[:80]}", flush=True)
                next_pending.append(item)
        if not next_pending:
            break
        if tries >= 4 and len(next_pending) > 1:
            split_total = len(next_pending)
            for idx, item in enumerate(next_pending, 1):
                label = f"{batch_label}|{idx}/{split_total}:{display_word(item['en'])}" if batch_label else display_word(item["en"])
                resolved.update(resolve_lexical_batch([item], label))
            return resolved
        if tries >= 8:
            raise RuntimeError(f"LLM 词汇关系校验失败次数过多: {batch_label or len(batch_items)}")
        pending = next_pending
    return resolved

def ensure_llm_lexical(entries):
    cache = load_llm_cache()
    missing = []
    seen = set()
    for en, cn in entries:
        key = material_key(en, cn)
        if key in seen:
            continue
        seen.add(key)
        if not has_lexical_fields(cache.get(key)):
            missing.append({"id": key, "en": en, "cn": cn})
    batch_size = int(os.getenv("VIVI_LLM_BATCH_SIZE", "60"))
    concurrency = max(1, int(os.getenv("VIVI_LLM_CONCURRENCY", "5")))
    batches = [missing[i:i + batch_size] for i in range(0, len(missing), batch_size)]
    total_batches = len(batches)
    if total_batches == 0:
        return cache
    workers = min(concurrency, total_batches)
    print(f"  [LLM词汇] 缺失 {len(missing)} 条，batch={batch_size}，并发={workers}", flush=True)
    if workers == 1:
        for idx, batch in enumerate(batches, 1):
            resolved = resolve_lexical_batch(batch, f"{idx}/{total_batches}")
            for key, value in resolved.items():
                cache[key] = {**cache.get(key, {}), **value}
            save_llm_cache()
        return cache
    with ThreadPoolExecutor(max_workers=workers) as executor:
        futures = {
            executor.submit(resolve_lexical_batch, batch, f"{idx}/{total_batches}"): idx
            for idx, batch in enumerate(batches, 1)
        }
        errors = []
        for future in as_completed(futures):
            try:
                resolved = future.result()
            except Exception as ex:
                errors.append(str(ex))
                continue
            for key, value in resolved.items():
                cache[key] = {**cache.get(key, {}), **value}
            save_llm_cache()
        if errors:
            raise RuntimeError(errors[0])
    return cache

class WordData:
    """缓存每个词的 WordNet 信息。关系候选只保留教学场景较可信的词。"""
    def __init__(self, en, cn, embed_syns=None, llm_lex=None):
        self.en = en
        self.cn = cn
        self.clean = clean_word(en)
        self.head = self.clean.split()[0] if self.clean else ""
        self.synsets = []
        self._def = ""
        self._syns = []
        self._ants = []
        self._embed_syns = []
        self._syns_common = []
        self._ants_common = []
        self._syns_all = []
        self._examples = []
        if llm_lex is not None:
            self._load_from_llm(llm_lex)
            return

        self.synsets = self._load_synsets()
        self._embed_syns = [normalize_relation_term(x) for x in (embed_syns or [])
                            if is_relation_term(x, self.clean or self.head)]
        if self.synsets:
            # 取与词性标记最接近的常用义
            s0 = self.synsets[0]
            self._def = s0.definition()
            # 收集候选同义词/反义词；过滤同根派生、词形变化和生僻格式。
            strict_syn_cands = {}
            broad_syn_cands = {}
            strict_ant_cands = {}
            broad_ant_cands = {}
            for idx, s in enumerate(self.synsets[:4]):
                for ex in s.examples():
                    if ex and ex not in self._examples:
                        self._examples.append(ex)
                for l in s.lemmas():
                    nm = normalize_relation_term(l.name())
                    allow_wordnet_syn = normalize_relation_term(self.clean) not in BAD_RELATION_HEADS
                    allow_wordnet_ant = normalize_relation_term(self.clean) not in BAD_ANTONYM_HEADS
                    if allow_wordnet_syn and nm not in BAD_RELATION_TERMS and is_relation_term(nm, self.clean or self.head):
                        score = max(commonness(nm.replace(' ', '_')), 1)
                        if idx == 0 and ' ' not in nm:
                            strict_syn_cands[nm] = max(strict_syn_cands.get(nm, 0), score)
                        if ' ' not in nm:
                            broad_syn_cands[nm] = max(broad_syn_cands.get(nm, 0), score)
                    for a in l.antonyms():
                        nm = normalize_relation_term(a.name())
                        if allow_wordnet_ant and nm not in BAD_RELATION_TERMS and is_relation_term(nm, self.clean or self.head):
                            score = max(commonness(nm.replace(' ', '_')), 1)
                            if idx <= 1:
                                strict_ant_cands[nm] = max(strict_ant_cands.get(nm, 0), score)
                            broad_ant_cands[nm] = max(broad_ant_cands.get(nm, 0), score)

            def ranked(cands):
                return [k for k, _ in sorted(cands.items(), key=lambda x: (-x[1], len(x[0]), x[0]))]

            curated_syns = curated_terms(CURATED_SYNONYMS, self.clean)
            curated_ants = curated_terms(CURATED_ANTONYMS, self.clean)
            self._syns_common = list(dict.fromkeys(curated_syns + ranked(strict_syn_cands)))
            self._ants_common = list(dict.fromkeys(curated_ants + ranked(strict_ant_cands)))
            self._syns = list(dict.fromkeys(self._syns_common + ranked(broad_syn_cands)))
            self._ants = list(dict.fromkeys(self._ants_common + ranked(broad_ant_cands)))
            # 嵌入层只作为补充，而且必须通过同样的关系词过滤。
            merged = []
            seen = set()
            for w in self._syns_common + self._syns + self._embed_syns:
                if w and w not in seen:
                    merged.append(w)
                    seen.add(w)
            self._syns_all = merged

        # Some phrases are not in WordNet; keep small hand-vetted overrides active.
        curated_syns = curated_terms(CURATED_SYNONYMS, self.clean)
        curated_ants = curated_terms(CURATED_ANTONYMS, self.clean)
        if curated_syns:
            self._syns_common = list(dict.fromkeys(curated_syns + self._syns_common))
            self._syns = list(dict.fromkeys(curated_syns + self._syns))
            self._syns_all = list(dict.fromkeys(curated_syns + self._syns_all))
        if curated_ants:
            self._ants_common = list(dict.fromkeys(curated_ants + self._ants_common))
            self._ants = list(dict.fromkeys(curated_ants + self._ants))

    def _load_from_llm(self, lex):
        self._def = str(lex.get("definition_en", "")).strip()
        syn = normalize_relation_term(lex.get("synonym", ""))
        ant = normalize_relation_term(lex.get("antonym", ""))
        if syn:
            self._syns_common = [syn]
            self._syns = [syn]
            self._syns_all = [syn]
        if ant:
            self._ants_common = [ant]
            self._ants = [ant]

    def _load_synsets(self):
        seen = set()
        out = []
        for key in wordnet_keys(self.en):
            for pos in pos_priority(self.cn):
                for syn in wn.synsets(key, pos=pos)[:4]:
                    if syn.name() not in seen:
                        out.append(syn)
                        seen.add(syn.name())
            if out:
                break
        return out
    @property
    def has_def(self): return bool(self._def)
    @property
    def defn(self): return self._def
    @property
    def syns(self): return self._syns_all or self._syns
    @property
    def syns_common(self):
        # 优先嵌入同义(已是常见词); 无嵌入时回退 WordNet 常见
        return self._embed_syns or self._syns_common or self._syns_all or self._syns
    @property
    def ants(self): return self._ants
    @property
    def ants_common(self): return self._ants_common or self._ants
    @property
    def examples(self): return self._examples
    @property
    def example(self):
        return self._examples[0] if self._examples else ""

def load_words(path, kind):
    """kind: 'xls' or 'xlsx'"""
    out = []
    if kind == 'xls':
        import xlrd
        wb = xlrd.open_workbook(path)
        ws = wb.sheet_by_index(0)
        for ri in range(2, ws.nrows):
            en = ws.cell_value(ri, 1)
            cn = ws.cell_value(ri, 3)
            if en:
                out.append((str(en).strip(), clean_cn(cn)))
    else:
        wb = openpyxl_loader(path)
        ws = wb["Sheet2"]
        for r in ws.iter_rows(min_row=3, values_only=True):
            if r[0] is not None and isinstance(r[0], (int, float)) and r[1]:
                out.append((str(r[1]).strip(), clean_cn(r[3]) if r[3] else ""))
    return out

def openpyxl_loader(path):
    import openpyxl
    return openpyxl.load_workbook(path, data_only=True)

# ============ 题型生成器 ============
# 每个生成器接收: (doc, answer_doc, group_words, group_idx, WData缓存dict)
# group_words: 本组50个 (en,cn) 元组
# 返回答案列表 [(num, letter_or_text)]

CN_NUM = "一二三四五六七八九十"

def pick_correct_index(n):
    return random.randint(0, n-1)

def make_distractors(pool, correct, n, key=lambda x:x):
    """从 pool 中排除 correct 后取 n 个干扰项"""
    cand = [x for x in pool if key(x) != key(correct)]
    random.shuffle(cand)
    return cand[:n]

def option_line(opts):
    return "    " + "  ".join(f"{chr(65+j)}. {o}" for j, o in enumerate(opts))

def fill_blank_sentence(en, cn, wd_item=None):
    """题型2用：优先用 WordNet 例句挖空，否则用稳定模板句。"""
    word = clean_word(en)
    if wd_item:
        for ex in wd_item.examples:
            pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
            if word and pattern.search(ex):
                return pattern.sub("______", ex, count=1)
    meaning = plain_cn(cn)
    return f"The word meaning \"{meaning}\" is ______."

def synonym_prompt(en, cn):
    word = display_word(en)
    meaning = plain_cn(cn)
    return (
        f"{word} means {meaning}.",
        f"—The word \"{word}\" can be replaced by _______.",
    )

BAD_SENTENCE_HEADWORDS = {
    "a", "between and", "cm", "good afternoon", "happy birthday", "how old", "kilo",
    "let's", "nothing", "to", "what about",
}

SENTENCE_STOPWORDS = {
    "a", "an", "and", "also", "are", "as", "at", "be", "but", "for", "have", "in",
    "is", "nor", "not", "of", "on", "only", "or", "sb", "sb's", "sth", "the", "to",
    "us", "what", "with", "your", "my", "his", "her", "our", "their", "one's",
}

def is_sentence_compatible_word(en):
    word = display_word(en)
    lower = word.lower()
    if lower in BAD_SENTENCE_HEADWORDS:
        return False
    if len(clean_word(word).replace(" ", "")) < 2:
        return False
    tokens = [tok for tok in re.findall(r"[A-Za-z']+", lower) if tok]
    content_tokens = [tok for tok in tokens if tok not in SENTENCE_STOPWORDS]
    if len(tokens) >= 2 and len(content_tokens) < 2:
        return False
    return bool(re.match(r"^[A-Za-z][A-Za-z' ]*[A-Za-z]$|^[A-Za-z]+$", word))

def ensure_pdf_font():
    global PDF_FONT_READY
    if PDF_FONT_READY:
        return
    pdfmetrics.registerFont(UnicodeCIDFont(TRANSLATION_PDF_FONT))
    PDF_FONT_READY = True

def pdf_text_width(text, font_name, font_size):
    ensure_pdf_font()
    return pdfmetrics.stringWidth(str(text), font_name, font_size)

def split_pdf_text(text, font_name, font_size, max_width):
    text = str(text).strip()
    if not text:
        return "", ""
    if pdf_text_width(text, font_name, font_size) <= max_width:
        return text, ""
    break_at = -1
    for idx in range(1, len(text) + 1):
        piece = text[:idx]
        if pdf_text_width(piece, font_name, font_size) > max_width:
            break
        ch = text[idx - 1]
        if ch.isspace() or ch in "，,。；;：:、/)-":
            break_at = idx
    if break_at <= 0:
        break_at = max(1, idx - 1)
    head = text[:break_at].rstrip()
    tail = text[break_at:].lstrip()
    if not head:
        head = text[:idx - 1]
        tail = text[idx - 1:].lstrip()
    return head, tail

def wrap_pdf_prefixed_line(prefix, text, font_name, font_size, max_width, base_indent=0):
    prefix = str(prefix)
    text = str(text).strip()
    prefix_width = pdf_text_width(prefix, font_name, font_size)
    first_width = max_width - base_indent - prefix_width
    next_indent = base_indent + prefix_width
    next_width = max_width - next_indent
    if first_width <= 8:
        first_width = max_width
    if next_width <= 8:
        next_width = max_width
    out = []
    remaining = text
    first = True
    while remaining:
        current_width = first_width if first else next_width
        head, tail = split_pdf_text(remaining, font_name, font_size, current_width)
        out.append({
            "indent": base_indent if first else next_indent,
            "text": (prefix + head) if first else head,
        })
        remaining = tail
        first = False
    if not out:
        out.append({"indent": base_indent, "text": prefix})
    return out

def build_translation_block(num, stem, opts, font_size, col_width):
    font_name = TRANSLATION_PDF_FONT
    line_height = font_size * 1.12
    option_indent = font_size * 0.8
    lines = []
    lines.extend(wrap_pdf_prefixed_line(f"{num}. ", stem, font_name, font_size, col_width, base_indent=0))
    for letter, opt in zip("ABCD", opts):
        lines.extend(wrap_pdf_prefixed_line(f"{letter}. ", opt, font_name, font_size, col_width, base_indent=option_indent))
    block_gap = max(2.0, font_size * 0.28)
    height = len(lines) * line_height + block_gap
    return {
        "num": num,
        "stem": stem,
        "opts": list(opts),
        "font_name": font_name,
        "font_size": font_size,
        "line_height": line_height,
        "block_gap": block_gap,
        "lines": lines,
        "height": height,
    }

def plan_translation_layout(rows, cols, font_size, spacing_pt):
    page_width, page_height = TRANSLATION_PDF_PAGE_SIZE
    usable_width = page_width - TRANSLATION_PDF_LEFT_PT - TRANSLATION_PDF_RIGHT_PT
    usable_height = page_height - TRANSLATION_PDF_TOP_PT - TRANSLATION_PDF_BOTTOM_PT
    col_width = (usable_width - spacing_pt * (cols - 1)) / cols
    if col_width <= 0:
        return None
    blocks = [build_translation_block(i, stem, opts, font_size, col_width) for i, (stem, opts) in enumerate(rows, 1)]
    placements = []
    col_idx = 0
    used_height = 0.0
    counts = [0] * cols
    for block in blocks:
        if block["height"] > usable_height:
            return None
        if used_height and used_height + block["height"] > usable_height + 0.1:
            col_idx += 1
            used_height = 0.0
        if col_idx >= cols:
            return None
        placements.append({"column": col_idx, "y_offset": used_height, "block": block})
        used_height += block["height"]
        counts[col_idx] += 1
    return {
        "cols": cols,
        "font_size": font_size,
        "spacing_pt": spacing_pt,
        "col_width": col_width,
        "counts": counts,
        "placements": placements,
    }

def choose_translation_layout(rows):
    scored = []
    for candidate in TRANSLATION_LAYOUT_CANDIDATES:
        planned = plan_translation_layout(
            rows,
            candidate["cols"],
            candidate["font_size"],
            candidate["spacing_pt"],
        )
        if planned:
            counts = [count for count in planned["counts"] if count > 0]
            used_cols = len(counts)
            min_count = min(counts) if counts else 0
            spread = (max(counts) - min(counts)) if counts else 99
            fully_used = int(used_cols == planned["cols"])
            balanced_floor = min(min_count, 7)
            score = (
                planned["font_size"],
                fully_used,
                balanced_floor,
                -spread,
                planned["cols"],
            )
            scored.append((score, planned))
    if scored:
        scored.sort(key=lambda item: item[0], reverse=True)
        return scored[0][1]
    return plan_translation_layout(rows, 3, 8.5, 12) or {
        "cols": 3,
        "font_size": 8.5,
        "spacing_pt": 12,
        "col_width": (
            (TRANSLATION_PDF_PAGE_SIZE[0] - TRANSLATION_PDF_LEFT_PT - TRANSLATION_PDF_RIGHT_PT - 24) / 3
        ),
        "counts": [0, 0, 0],
        "placements": [],
    }

def prepare_translation_section(doc, gi, layout, header_text=""):
    if gi == 0:
        section = doc.sections[-1]
    else:
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        section.header.is_linked_to_previous = False
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(90)
    section.right_margin = Pt(90)
    section.header_distance = Pt(42.5)
    set_section_columns(section, layout["cols"], spacing_pt=layout["spacing_pt"])
    if header_text:
        set_section_header(section, header_text, font=STANDARD_FONT, size=TRANSLATION_HEADER_PT)
    return section

def add_translation_question_block(doc, num, stem, opts, font_size):
    paragraphs = [
        add_para(doc, f"{num}. {stem}", size=font_size, space_before=0, space_after=0, line=1.0),
        add_para(doc, f"\tA. {opts[0]}", size=font_size, space_before=0, space_after=0, line=1.0),
        add_para(doc, f"\tB. {opts[1]}", size=font_size, space_before=0, space_after=0, line=1.0),
        add_para(doc, f"\tC. {opts[2]}", size=font_size, space_before=0, space_after=0, line=1.0),
        add_para(doc, f"\tD. {opts[3]}", size=font_size, space_before=0, space_after=0, line=1.0),
    ]
    for idx, para in enumerate(paragraphs):
        set_paragraph_keep(para, keep_next=(idx < len(paragraphs) - 1), keep_lines=True)

def render_translation_questions_pdf(pdf_path, pages):
    ensure_pdf_font()
    page_width, page_height = TRANSLATION_PDF_PAGE_SIZE
    pdf = canvas.Canvas(pdf_path, pagesize=TRANSLATION_PDF_PAGE_SIZE)
    top_y = page_height - TRANSLATION_PDF_TOP_PT
    header_y = page_height - 42.5
    for page in pages:
        if page.get("header_text"):
            pdf.setFont(TRANSLATION_PDF_HEADER_FONT, TRANSLATION_HEADER_PT)
            pdf.drawCentredString(page_width / 2.0, header_y, page["header_text"])
        layout = page["layout"]
        col_width = layout["col_width"]
        col_gap = layout["spacing_pt"]
        for placement in layout["placements"]:
            block = placement["block"]
            col_idx = placement["column"]
            x = TRANSLATION_PDF_LEFT_PT + col_idx * (col_width + col_gap)
            y = top_y - placement["y_offset"]
            pdf.setFont(block["font_name"], block["font_size"])
            for line in block["lines"]:
                pdf.drawString(x + line["indent"], y, line["text"])
                y -= block["line_height"]
        pdf.showPage()
    pdf.save()

def draw_answer_chunk(pdf, left_x, y, chunk):
    usable_width = TRANSLATION_PDF_PAGE_SIZE[0] - TRANSLATION_PDF_LEFT_PT - TRANSLATION_PDF_RIGHT_PT
    cell_width = usable_width / 10.0
    for idx, (num, val) in enumerate(chunk):
        pdf.drawString(left_x + idx * cell_width, y, f"{num}.{val}")

def render_translation_answers_pdf(pdf_path, title, groups):
    ensure_pdf_font()
    page_width, page_height = TRANSLATION_PDF_PAGE_SIZE
    pdf = canvas.Canvas(pdf_path, pagesize=TRANSLATION_PDF_PAGE_SIZE)
    top_y = page_height - TRANSLATION_PDF_TOP_PT
    left_x = TRANSLATION_PDF_LEFT_PT
    y = top_y

    def new_page(with_title=False):
        nonlocal y
        if with_title:
            pdf.setFont(TRANSLATION_PDF_FONT, TRANSLATION_ANSWER_TITLE_PT)
            pdf.drawString(left_x, y, title)
            y -= TRANSLATION_ANSWER_TITLE_PT * 1.5

    new_page(with_title=True)
    for group in groups:
        lines = [group["answers"][i:i + 10] for i in range(0, len(group["answers"]), 10)]
        needed = (
            TRANSLATION_ANSWER_GROUP_PT * 1.4 +
            len(lines) * (TRANSLATION_ANSWER_LINE_PT * 1.25) +
            8
        )
        if y - needed < TRANSLATION_PDF_BOTTOM_PT:
            pdf.showPage()
            y = top_y
        pdf.setFont(TRANSLATION_PDF_FONT, TRANSLATION_ANSWER_GROUP_PT)
        pdf.drawString(left_x, y, group["range_title"])
        y -= TRANSLATION_ANSWER_GROUP_PT * 1.35
        pdf.setFont(TRANSLATION_PDF_FONT, TRANSLATION_ANSWER_LINE_PT)
        for chunk in lines:
            draw_answer_chunk(pdf, left_x, y, chunk)
            y -= TRANSLATION_ANSWER_LINE_PT * 1.2
        y -= 8
    pdf.save()

def choose_synonym_words(words, wd):
    words = [w for w in words if is_sentence_compatible_word(w[0])] or list(words)
    pool = [w for w in words if preferred_synonym_for_word(wd[w])]
    if len(pool) < 30:
        extra = [w for w in words if wd[w].syns_common and w not in set(pool)]
        if len(pool) + len(extra) < 30:
            extra += [w for w in words if wd[w].syns and w not in set(pool) and w not in set(extra)]
        pool = pool + extra
    if len(pool) >= 30:
        return random.sample(pool, 30)
    chosen = pool[:]
    i = 0
    while chosen and len(chosen) < 30:
        chosen.append(pool[i % len(pool)])
        i += 1
    random.shuffle(chosen)
    return chosen

# ---- 题型1: 释义匹配 (英义→单词) 每50选30出30题 ----
def gen_matching(doc, ans_doc, words, gi, wd, ctx=None):
    title = "一. Matching Words with Definitions 单词释义匹配题"
    add_para(doc, title, size=GROUP_SUB_PT, bold=True, space_before=6, space_after=4)
    add_para(doc, "Match each definition with the correct word. 根据英文释义，从方框中选出正确单词。",
             size=BODY_PT, space_after=6)
    # 选30个有释义的词
    pool = [w for w in words if wd[w].has_def]
    if len(pool) < 30:
        pool = words[:]
    chosen = random.sample(pool, 30) if len(pool) >= 30 else pool[:]
    answers = []
    for i, (en, cn) in enumerate(chosen, 1):
        d = wd[(en, cn)]
        defn = d.defn
        # 干扰项: 同组其他词
        distractors = unique_distractors(words, (en, cn), 3, key=lambda x: clean_word(x[0]))
        opts = [display_word(dd[0]) for dd in distractors] + [display_word(en)]
        random.shuffle(opts)
        correct_letter = "ABCD"[opts.index(display_word(en))]
        add_para(doc, f"{i}. {defn}", size=BODY_PT)
        add_para(doc, option_line(opts), size=BODY_PT)
        answers.append((i, correct_letter))
    # 答案
    write_answer_block(ans_doc, gi, "一 释义匹配", answers, kind="letter")
    return answers

# ---- 题型2: 选择题 (中文释义选英文) 每50选30出30题 ----
def gen_multiple_choice(doc, ans_doc, words, gi, wd, ctx=None):
    title = "二. Multiple-Choice Questions 单词选择题"
    add_para(doc, title, size=GROUP_SUB_PT, bold=True, space_before=6, space_after=4)
    add_para(doc, "Choose the best word to complete each sentence. 选择最佳单词补全句子。",
             size=BODY_PT, space_after=6)
    pool = [w for w in words if is_sentence_compatible_word(w[0])] or list(words)
    chosen = ctx.get("chosen") if ctx and ctx.get("chosen") else (random.sample(pool, 30) if len(pool) >= 30 else pool[:])
    materials = ensure_llm_materials(chosen, wd, require_synonym=False)
    answers = []
    for i, (en, cn) in enumerate(chosen, 1):
        # 干扰项: 同组其他词的英文
        distractors = unique_distractors(pool, (en,cn), 3, key=lambda x: clean_word(x[0]))
        opts = [display_word(d[0]) for d in distractors] + [display_word(en)]
        random.shuffle(opts)
        correct_letter = "ABCD"[opts.index(display_word(en))]
        mat = materials[material_key(en, cn)]
        add_para(doc, f"{i}. {mat['cloze_sentence']}", size=BODY_PT)
        add_para(doc, option_line(opts), size=BODY_PT)
        answers.append((i, correct_letter))
    write_answer_block(ans_doc, gi, "二 选择题", answers, kind="letter")
    return answers

# ---- 题型3: 同义替换 每50选30出30题 ----
def gen_synonym_replace(doc, ans_doc, words, gi, wd, ctx=None):
    title = "三. Synonym Replacement 同义替换"
    add_para(doc, title, size=GROUP_SUB_PT, bold=True, space_before=6, space_after=4)
    add_para(doc, "Replace the underlined word with its synonym. 用同义词替换句中画线单词。",
             size=BODY_PT, space_after=6)
    # 优先使用可靠同义候选；不足时不再用词形变化硬凑。
    pool = [w for w in words if wd[w].syns_common]
    if len(pool) < 30:
        extra = [w for w in words if wd[w].syns and w not in set(pool)]
        pool = pool + extra
    chosen = ctx.get("chosen") if ctx and ctx.get("chosen") else choose_synonym_words(words, wd)
    materials = ensure_llm_materials(chosen, wd, require_synonym=True)
    answers = []
    for i, (en, cn) in enumerate(chosen, 1):
        d = wd[(en, cn)]
        mat = materials[material_key(en, cn)]
        syn = mat.get("synonym") or (d.syns_common[0] if d.syns_common else (d.syns[0] if d.syns else display_word(en)))
        line1 = mat.get("synonym_original") or synonym_prompt(en, cn)[0]
        line2 = mat.get("synonym_rewrite_blank") or synonym_prompt(en, cn)[1]
        add_para(doc, f"{i}. {line1}", size=BODY_PT)
        add_para(doc, line2, size=BODY_PT)
        # 4选项: 1个同义 + 3个干扰(其他词)
        distractors = unique_distractors(words, (en,cn), 3, key=lambda x: clean_word(x[0]))
        opts = [display_word(d2[0]) for d2 in distractors] + [syn]
        # 去重保护
        opts = list(dict.fromkeys(opts))
        while len(opts) < 4:
            extra = display_word(random.choice([w[0] for w in words if display_word(w[0]) not in opts]))
            opts.append(extra)
        random.shuffle(opts)
        correct_letter = "ABCD"[opts.index(syn)] if syn in opts else "A"
        add_para(doc, option_line(opts[:4]), size=BODY_PT)
        answers.append((i, correct_letter))
    write_answer_block(ans_doc, gi, "三 同义替换", answers, kind="letter")
    return answers

# ---- 题型4: 单词乱序拼写 每50选10出10题 ----
def gen_scramble(doc, ans_doc, words, gi, wd, ctx=None):
    title = "四. Word Scramble 单词乱序拼写题"
    add_para(doc, title, size=GROUP_SUB_PT, bold=True, space_before=6, space_after=4)
    add_para(doc, "Rearrange the letters to form correct words. 重新排列字母，拼写出正确单词。",
             size=BODY_PT, space_after=6)
    # 选10个长度4~10、纯字母的单词(排除短语)
    pool = [w for w in words if 4 <= len(spelling_core(w[0])) <= 10]
    chosen = random.sample(pool, 10) if len(pool) >= 10 else pool[:]
    answers = []
    for i, (en, cn) in enumerate(chosen, 1):
        core = spelling_core(en)
        letters = list(core)
        # 打乱：保证不等于原词
        for _ in range(10):
            random.shuffle(letters)
            if "".join(letters) != core:
                break
        scrambled = " ".join(letters)
        add_para(doc, f"{i}. {scrambled} → ____________  ({plain_cn(cn)})", size=BODY_PT)
        answers.append((i, display_word(en)))
    write_answer_block(ans_doc, gi, "四 乱序拼写", answers, kind="text")
    return answers

# ---- 题型5: 缺字母填空 每50选10出10题 ----
def gen_missing_letters(doc, ans_doc, words, gi, wd, ctx=None):
    title = "五. Missing Letters 缺字母填空"
    add_para(doc, title, size=GROUP_SUB_PT, bold=True, space_before=6, space_after=4)
    add_para(doc, "Fill in the missing letters and write the full word. 补全所缺字母，并写出完整单词。",
             size=BODY_PT, space_after=6)
    pool = [w for w in words if 5 <= len(spelling_core(w[0])) <= 12]
    chosen = random.sample(pool, 10) if len(pool) >= 10 else pool[:]
    answers = []
    for i, (en, cn) in enumerate(chosen, 1):
        core = spelling_core(en)
        n = len(core)
        # 隐藏约1/3字母，保证首字母可见
        hide_count = max(1, n // 3)
        idxs = random.sample(range(1, n), min(hide_count, n-1))
        chars = list(core)
        for k in idxs:
            chars[k] = "_"
        display = " ".join(chars)
        add_para(doc, f"{i}. {display}  ({plain_cn(cn)})", size=BODY_PT)
        answers.append((i, display_word(en)))
    write_answer_block(ans_doc, gi, "五 缺字母填空", answers, kind="text")
    return answers

# ---- 题型6: 同义反义词辨析 每50选30出30题 ----
def gen_syn_ant_judge(doc, ans_doc, words, gi, wd, ctx=None):
    title = "六. Synonym & Antonym 同义反义词辨析"
    add_para(doc, title, size=GROUP_SUB_PT, bold=True, space_before=6, space_after=4)
    add_para(doc, "Write S (synonym 同义) or A (antonym 反义) for each pair. 每组单词括号内填写S或A。",
             size=BODY_PT, space_after=6)
    syn_pairs = []
    ant_pairs = []
    for w in words:
        d = wd[w]
        if d.syns_common:
            syn_pairs.append((display_word(w[0]), d.syns_common[0], 'S'))
        if d.ants_common:
            ant_pairs.append((display_word(w[0]), d.ants_common[0], 'A'))
    random.shuffle(syn_pairs)
    random.shuffle(ant_pairs)
    pairs = ant_pairs[:12] + syn_pairs[:18]
    if len(pairs) < 30:
        rest = [p for p in ant_pairs[12:] + syn_pairs[18:] if p not in pairs]
        pairs += rest[:30-len(pairs)]
    if pairs and len(pairs) < 30:
        # 只复用可信关系，不再随机生成未知关系。
        i = 0
        while len(pairs) < 30:
            pairs.append(pairs[i % len(pairs)])
            i += 1
    pairs = pairs[:30]
    random.shuffle(pairs)
    answers = []
    for i, (left, right, rel) in enumerate(pairs, 1):
        add_para(doc, f"{i}. {left} & {right} ( )", size=BODY_PT)
        answers.append((i, rel))
    write_answer_block(ans_doc, gi, "六 同义反义辨析", answers, kind="text")
    return answers

def judge_rel(wd, a, b):
    da, db = wd[a], wd[b]
    if b[0] in da.syns or a[0] in db.syns:
        return "S"
    if b[0] in da.ants or a[0] in db.ants:
        return "A"
    return "S"  # 默认(教学场景多数配对为同义)

# ---- 题型7: 同义词匹配 每50出6组每组5词 ----
def _match_blocks(doc, ans_doc, gi, section_name, words, wd, rel_attr, target_groups=6):
    """通用匹配题：左侧只从当前50词中取，右侧为可信同/反义候选。"""
    pairs = []
    for w in words:
        rels = list(getattr(wd[w], rel_attr))
        for rel in rels[:1]:
            pairs.append((w, rel))
    random.shuffle(pairs)
    answers = []
    qn = 0
    used_pair_idx = 0
    if not pairs:
        write_answer_block(ans_doc, gi, section_name, answers, kind="text")
        return answers
    for _ in range(target_groups):
        block = []
        used_left = set()
        used_right = set()
        attempts = 0
        while len(block) < 5 and attempts < max(80, len(pairs) * 4):
            pair = pairs[used_pair_idx % len(pairs)]
            used_pair_idx += 1
            attempts += 1
            left_key = clean_word(pair[0][0])
            if left_key in used_left or pair[1] in used_right:
                continue
            block.append(pair)
            used_left.add(left_key)
            used_right.add(pair[1])
        # 极少数组反义词不足5个时，复用可信项补满版式。
        while len(block) < 5:
            pair = pairs[used_pair_idx % len(pairs)]
            used_pair_idx += 1
            block.append(pair)
        right_words = [right for _, right in block]
        random.shuffle(right_words)
        letters = "abcde"
        add_para(doc, "", size=BODY_PT)
        for j, (w, correct_right) in enumerate(block):
            correct_letter = letters[right_words.index(correct_right)] if correct_right in right_words else letters[0]
            add_para(doc, f"{j+1}. {display_word(w[0]):<20} {letters[j]}. {right_words[j]}", size=BODY_PT)
            qn += 1
            answers.append((qn, f"{j+1}-{correct_letter}"))
    write_answer_block(ans_doc, gi, section_name, answers, kind="text")
    return answers

def gen_synonym_matching(doc, ans_doc, words, gi, wd, ctx=None):
    title = "七. Synonym Matching 同义词匹配"
    add_para(doc, title, size=GROUP_SUB_PT, bold=True, space_before=6, space_after=4)
    add_para(doc, "Match each word with its synonym on the right. 将左侧单词与右侧同义词连线匹配。",
             size=BODY_PT, space_after=6)
    return _match_blocks(doc, ans_doc, gi, "七 同义词匹配", words, wd, "syns_common")

# ---- 题型8: 反义词匹配 每50出6组每组5词 ----
def gen_antonym_matching(doc, ans_doc, words, gi, wd, ctx=None):
    title = "八. Antonym Matching 反义词匹配"
    add_para(doc, title, size=GROUP_SUB_PT, bold=True, space_before=6, space_after=4)
    add_para(doc, "Match each word with its antonym on the right. 将左侧单词与右侧反义词连线匹配。",
             size=BODY_PT, space_after=6)
    return _match_blocks(doc, ans_doc, gi, "八 反义词匹配", words, wd, "ants_common")

# ---- 题型9: 判断正误 每50选10出10题 ----
def gen_true_false(doc, ans_doc, words, gi, wd, ctx=None):
    title = "九. True or False 判断正误"
    add_para(doc, title, size=GROUP_SUB_PT, bold=True, space_before=6, space_after=4)
    add_para(doc, "Write T (true) or F (false) for each statement. 判断下列句子的正误，正确填T，错误填F。",
             size=BODY_PT, space_after=6)
    pool = [w for w in words if is_sentence_compatible_word(w[0])] or list(words)
    chosen = ctx.get("chosen") if ctx and ctx.get("chosen") else (random.sample(pool, 10) if len(pool) >= 10 else pool[:])
    materials = ensure_llm_materials(chosen, wd, require_synonym=False)
    answers = []
    for i, (en, cn) in enumerate(chosen, 1):
        mat = materials[material_key(en, cn)]
        if random.random() < 0.5:
            sent = mat["tf_false"]
            ans = "F"
        else:
            sent = mat["tf_true"]
            ans = "T"
        add_para(doc, f"{i}. ( ) {sent}", size=BODY_PT)
        answers.append((i, ans))
    write_answer_block(ans_doc, gi, "九 判断正误", answers, kind="text")
    return answers

# ---- 题型10: 汉译英 (参照汉译英.docx格式) 每50选30出30题 ----
def gen_cn2en(doc, ans_doc, words, gi, wd, ctx=None):
    pool = list(words)
    chosen = random.sample(pool, 30) if len(pool) >= 30 else pool[:]
    rows = []
    answers = []
    for i, (en, cn) in enumerate(chosen, 1):
        distractors = unique_distractors(pool, (en,cn), 3, key=lambda x: clean_word(x[0]))
        opts = [display_word(d[0]) for d in distractors] + [display_word(en)]
        random.shuffle(opts)
        correct_letter = "ABCD"[opts.index(display_word(en))]
        rows.append((plain_cn(cn), opts))
        answers.append((i, correct_letter))
    layout = choose_translation_layout(rows)
    range_start, range_end = translation_group_range(gi, words, ctx)
    header_text = f"{range_start}~{range_end}"
    prepare_translation_section(doc, gi, layout, header_text=header_text)
    for i, (stem, opts) in enumerate(rows, 1):
        add_translation_question_block(doc, i, stem, opts, layout["font_size"])
    write_translation_answer_block(ans_doc, range_start, range_end, answers)
    if ctx and ctx.get("translation_pages") is not None:
        ctx["translation_pages"].append({
            "range_start": range_start,
            "range_end": range_end,
            "header_text": header_text,
            "rows": rows,
            "layout": layout,
        })
    if ctx and ctx.get("translation_answer_pages") is not None:
        ctx["translation_answer_pages"].append({
            "range_title": f"{range_start}～{range_end}答案",
            "answers": answers,
        })
    return answers

# ---- 题型11: 英译汉 (参照英译汉.docx格式) 每50选30出30题 ----
def gen_en2cn(doc, ans_doc, words, gi, wd, ctx=None):
    pool = list(words)
    chosen = random.sample(pool, 30) if len(pool) >= 30 else pool[:]
    rows = []
    answers = []
    for i, (en, cn) in enumerate(chosen, 1):
        distractors = unique_distractors(pool, (en,cn), 3, key=lambda x: plain_cn(x[1]))
        opts = [plain_cn(d[1]) for d in distractors] + [plain_cn(cn)]
        random.shuffle(opts)
        correct_letter = "ABCD"[opts.index(plain_cn(cn))]
        rows.append((display_word(en), opts))
        answers.append((i, correct_letter))
    layout = choose_translation_layout(rows)
    range_start, range_end = translation_group_range(gi, words, ctx)
    header_text = f"{range_start}~{range_end}"
    prepare_translation_section(doc, gi, layout, header_text=header_text)
    for i, (stem, opts) in enumerate(rows, 1):
        add_translation_question_block(doc, i, stem, opts, layout["font_size"])
    write_translation_answer_block(ans_doc, range_start, range_end, answers)
    if ctx and ctx.get("translation_pages") is not None:
        ctx["translation_pages"].append({
            "range_start": range_start,
            "range_end": range_end,
            "header_text": header_text,
            "rows": rows,
            "layout": layout,
        })
    if ctx and ctx.get("translation_answer_pages") is not None:
        ctx["translation_answer_pages"].append({
            "range_title": f"{range_start}～{range_end}答案",
            "answers": answers,
        })
    return answers

def translation_group_range(gi, words, ctx=None):
    if ctx and ctx.get("group_range"):
        return ctx["group_range"]
    start = gi * GROUP_SIZE + 1
    return start, start + len(words) - 1

def write_translation_answer_block(ans_doc, range_start, range_end, answers):
    add_para(ans_doc, f"{range_start}～{range_end}答案", size=TRANSLATION_ANSWER_GROUP_PT,
             bold=True, font=STANDARD_FONT, space_before=0, space_after=0, line=1.0)
    for start_idx in range(0, len(answers), 10):
        chunk = answers[start_idx:start_idx+10]
        line = "\t".join(f"{num}.{val}" for num, val in chunk)
        add_para(ans_doc, line, size=TRANSLATION_ANSWER_LINE_PT,
                 font=STANDARD_FONT, space_before=0, space_after=0, line=1.0)

# ============ 答案文档写入 ============
def write_answer_block(ans_doc, gi, section_name, answers, kind="letter", range_title=False):
    """每组答案一个标题 + 排版行"""
    # 标题
    if range_title:
        start = gi * GROUP_SIZE + 1
        end = (gi + 1) * GROUP_SIZE
        title = f"{start}～{end}答案"
    else:
        title = f"第{gi+1}组 {section_name} 答案"
    add_para(ans_doc, title, size=ANSWER_GROUP_PT, bold=True, space_before=8, space_after=4)
    # 答案行: 每行10个, 用制表符分隔, 仿参考文档
    per_line = 5 if section_name.startswith(("四", "五")) else 10
    for start in range(0, len(answers), per_line):
        chunk = answers[start:start+per_line]
        parts = []
        for num, val in chunk:
            if kind == "letter":
                parts.append(f"{num}.{val}")
            else:
                parts.append(f"{num}.{val}")
        line = "\t".join(parts)
        add_para(ans_doc, line, size=BODY_PT, space_after=2)

# ============ 主流程 ============
QUESTION_TYPES = [
    ("一_释义匹配", gen_matching),
    ("二_选择题", gen_multiple_choice),
    ("三_同义替换", gen_synonym_replace),
    ("四_乱序拼写", gen_scramble),
    ("五_缺字母填空", gen_missing_letters),
    ("六_同义反义辨析", gen_syn_ant_judge),
    ("七_同义词匹配", gen_synonym_matching),
    ("八_反义词匹配", gen_antonym_matching),
    ("九_判断正误", gen_true_false),
    ("十_汉译英", gen_cn2en),
    ("十一_英译汉", gen_en2cn),
]

TRANSLATION_TEMPLATE_TYPES = {"十_汉译英", "十一_英译汉"}
WD_REQUIRED_TYPES = {
    "一_释义匹配",
    "二_选择题",
    "三_同义替换",
    "六_同义反义辨析",
    "七_同义词匹配",
    "八_反义词匹配",
    "九_判断正误",
}
LLM_BASIC_TYPES = {"二_选择题", "九_判断正误"}
LLM_SYNONYM_TYPES = {"三_同义替换"}

GROUP_SIZE = 50

def chunk_groups(words, size=GROUP_SIZE):
    return [words[i:i+size] for i in range(0, len(words), size)]

def build_worddata_cache(words, embed_syns_map=None):
    """构建词数据缓存。当前直接使用 LLM 生成词义/同反义，不再依赖 WordNet。"""
    lexical_cache = ensure_llm_lexical(words)
    wd = {}
    for en, cn in words:
        wd[(en, cn)] = WordData(en, cn, llm_lex=lexical_cache.get(material_key(en, cn), {}))
    return wd

def build_llm_choice_plan(groups, wd_cache):
    plan = {}
    basic_entries = []
    synonym_entries = []
    for gi, gwords in enumerate(groups):
        compatible = [w for w in gwords if is_sentence_compatible_word(w[0])] or list(gwords)
        chosen = random.sample(compatible, 30) if len(compatible) >= 30 else compatible[:]
        plan[("二_选择题", gi)] = chosen
        basic_entries.extend(chosen)

        chosen = choose_synonym_words(gwords, wd_cache)
        plan[("三_同义替换", gi)] = chosen
        synonym_entries.extend(chosen)

        chosen = random.sample(compatible, 10) if len(compatible) >= 10 else compatible[:]
        plan[("九_判断正误", gi)] = chosen
        basic_entries.extend(chosen)
    return plan, basic_entries, synonym_entries

def generate_for_list(words, out_root, list_name, embed_syns_map=None, question_types=None):
    question_types = question_types or QUESTION_TYPES
    active_keys = {qkey for qkey, _ in question_types}
    groups = chunk_groups(words)
    print(f"[{list_name}] 共 {len(groups)} 组")
    wd_cache = build_worddata_cache(words, embed_syns_map) if active_keys & WD_REQUIRED_TYPES else {}
    llm_choice_plan = {}
    if active_keys & (LLM_BASIC_TYPES | LLM_SYNONYM_TYPES):
        llm_choice_plan, llm_basic_entries, llm_synonym_entries = build_llm_choice_plan(groups, wd_cache)
        if active_keys & LLM_BASIC_TYPES:
            print(f"[{list_name}] 预热 LLM 基础题面材料 {len({material_key(en, cn) for en, cn in llm_basic_entries})} 条")
            ensure_llm_materials(llm_basic_entries, wd_cache, require_synonym=False)
        if active_keys & LLM_SYNONYM_TYPES:
            print(f"[{list_name}] 预热 LLM 同义替换题面材料 {len({material_key(en, cn) for en, cn in llm_synonym_entries})} 条")
            ensure_llm_materials(llm_synonym_entries, wd_cache, require_synonym=True)
    # 每个题型一个文件夹
    for qkey, qfunc in question_types:
        folder = os.path.join(out_root, list_name, qkey)
        os.makedirs(folder, exist_ok=True)
        translation_pages = []
        translation_answer_pages = []
        # 题目文档 & 答案文档
        qdoc = new_doc(page_size="a4" if qkey in TRANSLATION_TEMPLATE_TYPES else "letter")
        adoc = new_doc(page_size="a4" if qkey in TRANSLATION_TEMPLATE_TYPES else "letter")
        if qkey not in TRANSLATION_TEMPLATE_TYPES:
            add_para(qdoc, f"单词练习 · {qkey.split('_',1)[1]} ({list_name})",
                     size=TITLE_PT, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
            add_para(adoc, f"{qkey.split('_',1)[1]}答案 ({list_name})",
                     size=ANSWER_TITLE_PT, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
        else:
            _set_doc_normal_style(qdoc, STANDARD_FONT, TRANSLATION_NORMAL_PT)
            _set_doc_normal_style(adoc, STANDARD_FONT, TRANSLATION_NORMAL_PT)
            answer_title = "汉译英答案" if qkey == "十_汉译英" else "英译汉答案"
            add_para(adoc, answer_title, size=TRANSLATION_ANSWER_TITLE_PT, bold=True,
                     font=CN_FONT, space_before=0, space_after=0, line=1.0)
        for gi, gwords in enumerate(groups):
            # 组标题(题目文档)
            start = gi*GROUP_SIZE + 1
            end = min((gi+1)*GROUP_SIZE, len(words))
            if qkey not in TRANSLATION_TEMPLATE_TYPES:
                add_para(qdoc, f"第 {gi+1} 组（第{start}～{end}词）", size=QHEAD_PT, bold=True,
                         space_before=10, space_after=4)
            ctx = {
                "chosen": llm_choice_plan.get((qkey, gi)),
                "group_range": (start, end),
            }
            if qkey in TRANSLATION_TEMPLATE_TYPES:
                ctx["translation_pages"] = translation_pages
                ctx["translation_answer_pages"] = translation_answer_pages
            qfunc(qdoc, adoc, gwords, gi, wd_cache, ctx)
            if qkey not in TRANSLATION_TEMPLATE_TYPES and gi < len(groups)-1:
                qdoc.add_page_break()
        qpath = os.path.join(folder, f"{qkey}.docx")
        apath = os.path.join(folder, f"{qkey}答案.docx")
        if qkey in TRANSLATION_TEMPLATE_TYPES:
            qpdf = os.path.join(folder, f"{qkey}.pdf")
            apdf = os.path.join(folder, f"{qkey}答案.pdf")
            answer_title = "汉译英答案" if qkey == "十_汉译英" else "英译汉答案"
            render_translation_questions_pdf(qpdf, translation_pages)
            render_translation_answers_pdf(apdf, answer_title, translation_answer_pages)
            for stale_path in (qpath, apath):
                if os.path.exists(stale_path):
                    os.remove(stale_path)
            print(f"  ✓ {qkey}: {qpdf}")
        else:
            qdoc.save(qpath)
            adoc.save(apath)
            print(f"  ✓ {qkey}: {qpath}")

def main():
    base = "/Users/boooo/vivi"
    out_root = os.path.join(base, "成品")
    # 这版不用嵌入近邻补同义词；它会把词形变化误当同义词。
    emb1350 = {}
    emb2200 = {}
    words1350 = load_words(os.path.join(base, "按难度排序词汇表.xls"), 'xls')
    words2200 = load_words(os.path.join(base, "2022年人教版中考英语单词汇总乱序版默写版.xlsx"), 'xlsx')
    # 1350 词表 (.xls)
    print(f"1350词表: {len(words1350)} 词")
    generate_for_list(words1350, out_root, "1350词组", emb1350)
    # 2200 词表 (.xlsx)
    print(f"2200词表: {len(words2200)} 词")
    generate_for_list(words2200, out_root, "2200词组", emb2200)
    print("\n全部生成完毕。输出目录:", out_root)

if __name__ == "__main__":
    main()
